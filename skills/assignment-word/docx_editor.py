"""
docx_editor.py — Edit and fill Word (.docx) documents.

Commands:
    info     — Print document structure (paragraphs, tables, placeholders)
    replace  — Find and replace text (preserving formatting)
    fill     — Batch replace placeholders from JSON
    image    — Insert image after matching text
    table    — Fill table data from JSON

Usage:
    python docx_editor.py info <docx>
    python docx_editor.py replace <docx> --old X --new Y [-o out.docx]
    python docx_editor.py fill <docx> --data replacements.json [-o out.docx]
    python docx_editor.py image <docx> --after "图1" --img fig.png [-o out.docx]
    python docx_editor.py table <docx> --idx 0 --data table.json [-o out.docx]
"""

import sys
import os
import json
import shutil
import argparse

try:
    from docx import Document
    from docx.shared import Cm, Inches, Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not installed. Run: python -m pip install python-docx")
    sys.exit(1)

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# Default Chinese font
DEFAULT_FONT = '宋体'


# ============================================================
# Font Helpers (East Asian font fix)
# ============================================================

def set_run_font(run, font_name=DEFAULT_FONT):
    """Set both Western and East Asian font on a run.

    python-docx's run.font.name only sets ascii/hAnsi (Western fonts).
    Chinese characters use w:eastAsia which must be set via XML,
    otherwise they render in Word's default font (e.g. Calibri) causing garbled text.
    """
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def set_style_font(style, font_name=DEFAULT_FONT):
    """Set both Western and East Asian font on a document style."""
    style.font.name = font_name
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


# ============================================================
# XML Helpers (from md2docx.py)
# ============================================================

def get_tcs(table, row_idx):
    """Get all actual w:tc elements for a row."""
    tr = table.rows[row_idx]._tr
    return tr.findall(f'{{{W}}}tc')


def tc_info(tc):
    """Get gridSpan and vMerge info for a tc element."""
    tcPr = tc.find(f'{{{W}}}tcPr')
    span = 1
    vmerge = False
    vmerge_restart = False
    if tcPr is not None:
        gs = tcPr.find(f'{{{W}}}gridSpan')
        if gs is not None:
            span = int(gs.get(f'{{{W}}}val', '1'))
        vm = tcPr.find(f'{{{W}}}vMerge')
        if vm is not None:
            vmerge = True
            val = vm.get(f'{{{W}}}val')
            vmerge_restart = val == 'restart'
    return span, vmerge, vmerge_restart


def set_tc_text(tc, text):
    """Set text content of a tc element (clears existing content). Sets East Asian font."""
    ps = tc.findall(f'{{{W}}}p')
    for p in ps[1:]:
        tc.remove(p)
    p0 = ps[0] if ps else tc.makeelement(f'{{{W}}}p', {})
    if not ps:
        tc.append(p0)
    for r in list(p0.findall(f'{{{W}}}r')):
        p0.remove(r)
    for d in list(p0.findall(f'{{{W}}}drawing')):
        p0.remove(d)
    new_r = p0.makeelement(f'{{{W}}}r', {})
    # Set East Asian font on the new run
    rPr = new_r.makeelement(f'{{{W}}}rPr', {})
    rFonts = rPr.makeelement(f'{{{W}}}rFonts', {
        qn('w:eastAsia'): DEFAULT_FONT,
        qn('w:ascii'): DEFAULT_FONT,
        qn('w:hAnsi'): DEFAULT_FONT,
    })
    rPr.append(rFonts)
    new_r.append(rPr)
    new_t = new_r.makeelement(f'{{{W}}}t', {})
    new_t.text = text
    new_r.append(new_t)
    p0.append(new_r)


def get_tc_text(tc):
    """Get text content of a tc element."""
    texts = [t.text for t in tc.iter(f'{{{W}}}t') if t.text]
    return ''.join(texts).strip()


# ============================================================
# Core Functions
# ============================================================

def replace_in_paragraph(paragraph, old_text, new_text):
    """Replace text in a paragraph, preserving run formatting."""
    full_text = paragraph.text
    if old_text not in full_text:
        return False

    # Strategy: find which runs contain the old text and replace
    # Build a mapping of character position to run
    runs = paragraph.runs
    if not runs:
        return False

    # Build char->run mapping
    char_run_map = []  # [(char, run_index)]
    for ri, run in enumerate(runs):
        for ch in run.text:
            char_run_map.append((ch, ri))

    # Find all occurrences
    text = ''.join(c for c, _ in char_run_map)
    occurrences = []
    start = 0
    while True:
        idx = text.find(old_text, start)
        if idx == -1:
            break
        occurrences.append(idx)
        start = idx + 1

    if not occurrences:
        return False

    # For simplicity, handle the case where old_text is within a single run
    # or spans multiple runs
    replaced = False
    for occ_start in reversed(occurrences):  # reverse to preserve indices
        occ_end = occ_start + len(old_text)

        # Find runs involved
        start_run = char_run_map[occ_start][1]
        end_run = char_run_map[occ_end - 1][1]

        if start_run == end_run:
            # Simple case: all within one run
            run = runs[start_run]
            run_text = run.text
            # Find position within this run
            run_char_start = sum(1 for c, ri in char_run_map[:occ_start] if ri == start_run)
            run_char_end = run_char_start + len(old_text)
            run.text = run_text[:run_char_start] + new_text + run_text[run_char_end:]
            replaced = True
        else:
            # Multi-run replacement
            # Put replacement text in the first run, clear the rest
            start_char_in_run = sum(1 for c, ri in char_run_map[:occ_start] if ri == start_run)
            runs[start_run].text = runs[start_run].text[:start_char_in_run] + new_text

            # Clear text in middle runs
            for ri in range(start_run + 1, end_run):
                runs[ri].text = ""

            # Remove the portion from the last run
            end_char_in_run = sum(1 for c, ri in char_run_map[:occ_end] if ri == end_run)
            runs[end_run].text = runs[end_run].text[end_char_in_run:]
            replaced = True

    return replaced


def replace_in_doc(doc, old_text, new_text):
    """Replace text throughout the document, preserving formatting."""
    count = 0

    # Replace in body paragraphs
    for para in doc.paragraphs:
        if replace_in_paragraph(para, old_text, new_text):
            count += 1

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if replace_in_paragraph(para, old_text, new_text):
                        count += 1

    # Replace in headers and footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header and header.is_linked_to_previous is False:
                for para in header.paragraphs:
                    if replace_in_paragraph(para, old_text, new_text):
                        count += 1
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer and footer.is_linked_to_previous is False:
                for para in footer.paragraphs:
                    if replace_in_paragraph(para, old_text, new_text):
                        count += 1

    return count


def insert_image_after_text(doc, search_text, img_path, width_cm=10.0):
    """Insert an image in a new paragraph after the paragraph containing search_text."""
    if not os.path.exists(img_path):
        print(f"ERROR: Image not found: {img_path}")
        return False

    # Search in body paragraphs
    for i, para in enumerate(doc.paragraphs):
        if search_text in para.text:
            # Insert a new paragraph after this one
            new_para = doc.add_paragraph()
            # Move the new paragraph to right after the current one
            para._element.addnext(new_para._element)
            run = new_para.add_run()
            run.add_picture(img_path, width=Cm(width_cm))
            print(f"Inserted image after paragraph {i}: '{para.text[:50]}...'")
            return True

    # Search in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if search_text in para.text:
                        new_para = cell.add_paragraph()
                        run = new_para.add_run()
                        run.add_picture(img_path, width=Cm(width_cm))
                        print(f"Inserted image in table cell: '{para.text[:50]}...'")
                        return True

    print(f"WARNING: Text '{search_text}' not found in document")
    return False


def fill_table_from_json(doc, table_idx, data):
    """Fill a table with data from JSON. Data format: {"cells": {"row,col": "value"}}"""
    if table_idx >= len(doc.tables):
        print(f"ERROR: Table index {table_idx} out of range (total: {len(doc.tables)})")
        return 0

    table = doc.tables[table_idx]
    cells_data = data.get("cells", {})
    count = 0

    for key, value in cells_data.items():
        parts = key.split(",")
        if len(parts) != 2:
            continue
        row_idx, col_idx = int(parts[0]), int(parts[1])

        if row_idx >= len(table.rows):
            print(f"  WARNING: Row {row_idx} out of range")
            continue

        tcs = get_tcs(table, row_idx)
        if col_idx >= len(tcs):
            print(f"  WARNING: Col {col_idx} out of range in row {row_idx}")
            continue

        tc = tcs[col_idx]
        _, vmerge, _ = tc_info(tc)
        if vmerge and not vmerge_restart:
            # Skip vertically merged continuation cells
            continue

        set_tc_text(tc, str(value))
        count += 1
        print(f"  Filled [{row_idx},{col_idx}]: {str(value)[:40]}")

    return count


# ============================================================
# Command: info
# ============================================================

def cmd_info(docx_path):
    """Print document structure."""
    doc = Document(docx_path)
    print(f"Document: {docx_path}")
    print(f"Sections: {len(doc.sections)}")
    print(f"Paragraphs: {len(doc.paragraphs)}")
    print(f"Tables: {len(doc.tables)}")
    print()

    print("=" * 60)
    print("PARAGRAPHS")
    print("=" * 60)
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            print(f"  [{i:3d}] <empty>")
            continue
        style = para.style.name if para.style else "None"
        # Truncate long text
        display = text[:80] + "..." if len(text) > 80 else text
        print(f"  [{i:3d}] [{style}] {display}")

    print()
    print("=" * 60)
    print("TABLES")
    print("=" * 60)
    for ti, table in enumerate(doc.tables):
        print(f"\n  Table {ti}: {len(table.rows)} rows x {len(table.columns)} cols")
        for ri in range(min(len(table.rows), 10)):  # Show first 10 rows
            tcs = get_tcs(table, ri)
            cells_text = []
            for ci, tc in enumerate(tcs):
                t = get_tc_text(tc)
                span, vm, vmr = tc_info(tc)
                suffix = ""
                if span > 1:
                    suffix += f"[span={span}]"
                if vm:
                    suffix += "[vmerge]"
                t = t[:30] + ".." if len(t) > 30 else t
                cells_text.append(f"{t}{suffix}")
            print(f"    Row {ri}: {' | '.join(cells_text)}")
        if len(table.rows) > 10:
            print(f"    ... ({len(table.rows) - 10} more rows)")

    # Headers/Footers
    print()
    print("=" * 60)
    print("HEADERS / FOOTERS")
    print("=" * 60)
    for si, section in enumerate(doc.sections):
        for name, part in [("Header", section.header),
                           ("Footer", section.footer)]:
            if part and part.is_linked_to_previous is False:
                for pi, para in enumerate(part.paragraphs):
                    text = para.text.strip()
                    if text:
                        print(f"  Section {si} {name} [{pi}]: {text[:60]}")


# ============================================================
# Main CLI
# ============================================================

def main():
    parser = argparse.ArgumentParser(description="Word document editor")
    subparsers = parser.add_subparsers(dest="command", help="Command to execute")

    # info
    p_info = subparsers.add_parser("info", help="Print document structure")
    p_info.add_argument("docx", help="Input .docx file")

    # replace
    p_replace = subparsers.add_parser("replace", help="Find and replace text")
    p_replace.add_argument("docx", help="Input .docx file")
    p_replace.add_argument("--old", required=True, help="Text to find")
    p_replace.add_argument("--new", required=True, help="Replacement text")
    p_replace.add_argument("-o", "--output", help="Output file path")

    # fill
    p_fill = subparsers.add_parser("fill", help="Batch fill placeholders from JSON")
    p_fill.add_argument("docx", help="Input .docx file")
    p_fill.add_argument("--data", required=True, help="JSON file with replacements")
    p_fill.add_argument("-o", "--output", help="Output file path")

    # image
    p_image = subparsers.add_parser("image", help="Insert image after text")
    p_image.add_argument("docx", help="Input .docx file")
    p_image.add_argument("--after", required=True, help="Text to search for")
    p_image.add_argument("--img", required=True, help="Image file path")
    p_image.add_argument("-o", "--output", help="Output file path")
    p_image.add_argument("--width", type=float, default=10.0, help="Image width in cm")

    # table
    p_table = subparsers.add_parser("table", help="Fill table from JSON")
    p_table.add_argument("docx", help="Input .docx file")
    p_table.add_argument("--idx", type=int, required=True, help="Table index (0-based)")
    p_table.add_argument("--data", required=True, help="JSON file with table data")
    p_table.add_argument("-o", "--output", help="Output file path")

    args = parser.parse_args()

    if not args.command:
        parser.print_help()
        sys.exit(1)

    if args.command == "info":
        cmd_info(args.docx)
        return

    # For edit commands, handle output path and backup
    docx_path = args.docx
    output_path = getattr(args, 'output', None)

    if output_path is None:
        # In-place edit with backup
        backup_path = docx_path + ".bak"
        shutil.copy2(docx_path, backup_path)
        print(f"Backup: {backup_path}")
        output_path = docx_path

    doc = Document(docx_path)

    if args.command == "replace":
        count = replace_in_doc(doc, args.old, args.new)
        print(f"Replaced '{args.old}' -> '{args.new}': {count} occurrences")

    elif args.command == "fill":
        with open(args.data, 'r', encoding='utf-8') as f:
            replacements = json.load(f)
        total = 0
        for old_text, new_text in replacements.items():
            count = replace_in_doc(doc, old_text, str(new_text))
            total += count
            if count > 0:
                print(f"  '{old_text}' -> '{str(new_text)[:50]}': {count} replacements")
            else:
                print(f"  '{old_text}': NOT FOUND")
        print(f"Total replacements: {total}")

    elif args.command == "image":
        insert_image_after_text(doc, args.after, args.img, args.width)

    elif args.command == "table":
        with open(args.data, 'r', encoding='utf-8') as f:
            table_data = json.load(f)
        count = fill_table_from_json(doc, args.idx, table_data)
        print(f"Filled {count} cells in table {args.idx}")

    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    main()
