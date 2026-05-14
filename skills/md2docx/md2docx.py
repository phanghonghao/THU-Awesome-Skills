"""
md2docx.py — Read Markdown data and fill into an existing Word (.docx) document.

Usage:
    python md2docx.py <input.md> <template.docx> [output.docx]

Handles:
- Table data matching by header text
- Merged cells (horizontal gridSpan, vertical vMerge)
- Image insertion from relative paths
"""

import sys
import os
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Cm, Inches, Pt
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed. Run: python -m pip install python-docx")
    sys.exit(1)

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


# ============================================================
# XML Helpers
# ============================================================

def get_tcs(table, row_idx):
    """Get all actual w:tc elements for a row (bypasses python-docx merge mapping)."""
    tr = table.rows[row_idx]._tr
    return tr.findall(f'{{{W}}}tc')


def tc_info(tc):
    """Get gridSpan and vMerge info for a tc element."""
    tcPr = tc.find(f'{{{W}}}tcPr')
    span = 1
    vmerge = False
    vmerge_restart = False
    if tcPr is not None:
        gs = tcPr.find(f'{{{W}}}gridSpan')
        if gs is not None:
            span = int(gs.get(f'{{{W}}}val', '1'))
        vm = tcPr.find(f'{{{W}}}vMerge')
        if vm is not None:
            vmerge = True
            val = vm.get(f'{{{W}}}val')
            vmerge_restart = val == 'restart'
    return span, vmerge, vmerge_restart


def set_tc_text(tc, text):
    """Set text content of a tc element (clears existing content)."""
    # Remove extra paragraphs
    ps = tc.findall(f'{{{W}}}p')
    for p in ps[1:]:
        tc.remove(p)
    # Clear first paragraph's runs
    p0 = ps[0] if ps else tc.makeelement(f'{{{W}}}p', {})
    if not ps:
        tc.append(p0)
    for r in list(p0.findall(f'{{{W}}}r')):
        p0.remove(r)
    # Also remove any drawing elements
    for d in list(p0.findall(f'{{{W}}}drawing')):
        p0.remove(d)
    # Create new run with text
    new_r = p0.makeelement(f'{{{W}}}r', {})
    new_t = new_r.makeelement(f'{{{W}}}t', {})
    new_t.text = text
    new_r.append(new_t)
    p0.append(new_r)


def get_tc_text(tc):
    """Get text content of a tc element."""
    texts = [t.text for t in tc.iter(f'{{{W}}}t') if t.text]
    return ''.join(texts).strip()


def set_tc_image(tc, img_path, width_cm=4.0, table=None, row_idx=None):
    """Insert image into a tc element. Needs table/row_idx to find Cell wrapper."""
    if not os.path.exists(img_path):
        print(f"  WARNING: Image not found: {img_path}")
        return False

    # Clear existing content
    ps = tc.findall(f'{{{W}}}p')
    for p in ps[1:]:
        tc.remove(p)
    p0 = ps[0]
    for r in list(p0.findall(f'{{{W}}}r')):
        p0.remove(r)

    # Find Cell wrapper for this tc
    if table is not None and row_idx is not None:
        row = table.rows[row_idx]
        for cell in row.cells:
            if cell._tc is tc:
                run = cell.paragraphs[0].add_run()
                run.add_picture(img_path, width=Cm(width_cm))
                return True

    print(f"  WARNING: Could not find Cell wrapper for image insertion")
    return False


# ============================================================
# Markdown Parser
# ============================================================

def parse_md_tables(md_text):
    """Parse all tables from Markdown text. Returns list of (headers, rows)."""
    tables = []
    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('|') and '|' in line[1:]:
            # Start of a table
            header_line = line
            rows = []
            i += 1
            # Skip separator line
            if i < len(lines) and '---' in lines[i]:
                i += 1
            # Read data rows
            while i < len(lines) and lines[i].strip().startswith('|'):
                row_line = lines[i].strip()
                if '---' in row_line:
                    i += 1
                    continue
                cells = [c.strip() for c in row_line.split('|')[1:-1]]
                rows.append(cells)
                i += 1
            headers = [c.strip() for c in header_line.split('|')[1:-1]]
            tables.append((headers, rows))
        else:
            i += 1
    return tables


def find_md_images(md_text, md_dir):
    """Find all image references in Markdown. Returns dict of {label: path}."""
    images = {}
    # Match ![label](path)
    pattern = r'!\[([^\]]*)\]\(([^)]+)\)'
    for match in re.finditer(pattern, md_text):
        label = match.group(1)
        path = match.group(2)
        full_path = os.path.join(md_dir, path)
        images[label] = full_path
        images[path] = full_path
    return images


# ============================================================
# Table Matching
# ============================================================

def match_md_to_docx_table(md_headers, md_rows, docx_table):
    """Try to match a Markdown table to a docx table by comparing headers.

    Returns (match_score, header_mapping) where header_mapping maps
    md column index → docx tc index.
    """
    # Get docx header row (find first non-empty row)
    best_score = 0
    best_mapping = None

    for ri in range(min(5, len(docx_table.rows))):
        tcs = get_tcs(docx_table, ri)
        docx_headers = []
        tc_to_grid = {}
        grid = 0
        for ti, tc in enumerate(tcs):
            span, vm, vmr = tc_info(tc)
            text = get_tc_text(tc)
            docx_headers.append((ti, grid, text))
            tc_to_grid[ti] = grid
            grid += span

        # Try to match md headers to docx headers
        matches = 0
        mapping = {}
        for mi, mh in enumerate(md_headers):
            mh_clean = re.sub(r'<[^>]+>', '', mh).strip().lower()
            for ti, g, dh in docx_headers:
                dh_clean = re.sub(r'<[^>]+>', '', dh).strip().lower()
                # Check for partial match
                if mh_clean and dh_clean and (mh_clean in dh_clean or dh_clean in mh_clean):
                    if len(mh_clean) > 1:  # Skip very short matches
                        matches += 1
                        mapping[mi] = ti
                        break

        if matches > best_score:
            best_score = matches
            best_mapping = mapping

    return best_score, best_mapping


# ============================================================
# Main Fill Logic
# ============================================================

def fill_docx(md_path, docx_path, out_path=None):
    """Read Markdown and fill data into the Word document."""

    if out_path is None:
        base, ext = os.path.splitext(docx_path)
        out_path = f"{base}_filled{ext}"

    # Read Markdown
    md_dir = os.path.dirname(os.path.abspath(md_path))
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    md_tables = parse_md_tables(md_text)
    md_images = find_md_images(md_text, md_dir)

    print(f"Parsed {len(md_tables)} tables and {len(md_images)} images from Markdown")

    # Open Word document
    doc = Document(docx_path)
    print(f"Opened docx with {len(doc.tables)} tables")

    # For each docx table, try to match and fill
    for dti, docx_table in enumerate(doc.tables):
        best_score = 0
        best_mti = -1
        best_mapping = None

        for mti, (md_headers, md_rows) in enumerate(md_tables):
            score, mapping = match_md_to_docx_table(md_headers, md_rows, docx_table)
            if score > best_score:
                best_score = score
                best_mti = mti
                best_mapping = mapping

        if best_score < 2:
            continue  # No good match

        md_headers, md_rows = md_tables[best_mti]
        print(f"\nTable {dti}: matched with MD table {best_mti} (score={best_score}, mapping={best_mapping})")

        # Analyze docx table structure
        num_rows = len(docx_table.rows)
        # Find header row (usually one with most text)
        header_row = 0
        max_text = 0
        for ri in range(min(5, num_rows)):
            tcs = get_tcs(docx_table, ri)
            total = sum(len(get_tc_text(tc)) for tc in tcs)
            if total > max_text:
                max_text = total
                header_row = ri

        # Find first data row (after header)
        data_start = header_row + 1
        # Skip separator rows (rows with ---)
        while data_start < num_rows:
            tcs = get_tcs(docx_table, data_start)
            text = ''.join(get_tc_text(tc) for tc in tcs)
            if '---' in text:
                data_start += 1
            else:
                break

        # Fill data rows
        for mri, md_row in enumerate(md_rows):
            dri = data_start + mri
            if dri >= num_rows:
                break

            tcs = get_tcs(docx_table, dri)

            for md_ci, value in enumerate(md_row):
                if not value or value.strip() == '':
                    continue

                if md_ci not in best_mapping:
                    continue

                tc_idx = best_mapping[md_ci]
                if tc_idx >= len(tcs):
                    continue

                tc = tcs[tc_idx]
                _, vmerge, _ = tc_info(tc)

                # Check if it's an image reference
                img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', value)
                if img_match:
                    img_path = os.path.join(md_dir, img_match.group(2))
                    set_tc_image(tc, img_path, width_cm=4.0, table=docx_table, row_idx=dri)
                else:
                    # Check if value looks like a reference to an image (e.g., "图3-1")
                    set_tc_text(tc, value.strip())

            print(f"  Filled row {dri}")

    # Save
    doc.save(out_path)
    print(f"\nSaved to: {out_path}")
    return out_path


# ============================================================
# CLI
# ============================================================

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python md2docx.py <input.md> <template.docx> [output.docx]")
        sys.exit(1)

    md_path = sys.argv[1]
    docx_path = sys.argv[2]
    out_path = sys.argv[3] if len(sys.argv) > 3 else None

    fill_docx(md_path, docx_path, out_path)
