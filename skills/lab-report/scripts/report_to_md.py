#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
实验报告 → MD 模板 提取器
从 PDF/DOCX 提取报告结构，生成带占位符的 Markdown 模板
复用 assignment 的 md_to_latex.py 进行 MD→TEX→PDF 编译
"""

import re
import os
import sys
import json
import shutil
import subprocess
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Tuple, Optional

# ── 依赖检测 ──
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    from docx import Document as DocxDocument
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

try:
    import mammoth
    HAS_MAMMOTH = True
except ImportError:
    HAS_MAMMOTH = False

# ── 复用 assignment 的 MD→TEX→PDF ──
ASSIGNMENT_DIR = str(Path(__file__).resolve().parent.parent.parent / "assignment")


def _import_assignment():
    """导入 assignment 的 md_to_latex 模块"""
    if ASSIGNMENT_DIR not in sys.path:
        sys.path.insert(0, ASSIGNMENT_DIR)
    try:
        import md_to_latex as m2l
        return m2l
    except ImportError:
        print(f"警告: 无法导入 assignment/md_to_latex.py (路径: {ASSIGNMENT_DIR})")
        return None


# ============================================================================
# PDF 提取
# ============================================================================

def extract_pdf(file_path: str) -> dict:
    """从 PDF 提取文字和图片"""
    if not HAS_PDFPLUMBER:
        print("错误: 需要安装 pdfplumber: pip install pdfplumber")
        return {"text": "", "tables": [], "images": []}

    text_parts = []
    tables = []

    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            # 提取文字
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

            # 提取表格
            page_tables = page.extract_tables()
            for t in page_tables:
                tables.append({
                    "page": i + 1,
                    "data": t
                })

    full_text = "\n\n".join(text_parts)
    return {
        "text": full_text,
        "tables": tables,
        "source": "pdf"
    }


# ============================================================================
# DOCX 提取
# ============================================================================

def extract_docx(file_path: str) -> dict:
    """从 DOCX 提取文字和表格"""
    if not HAS_PYTHON_DOCX:
        print("错误: 需要安装 python-docx: pip install python-docx")
        return {"text": "", "tables": []}

    doc = DocxDocument(file_path)

    text_parts = []
    tables = []

    # 提取段落文字
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            style = para.style.name if para.style else ""
            text_parts.append({"text": text, "style": style})

    # 提取表格
    for i, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)

    return {
        "paragraphs": text_parts,
        "tables": tables,
        "source": "docx"
    }


# ============================================================================
# 报告结构识别
# ============================================================================

# 章节关键词
SECTION_KEYWORDS = {
    "实验目的": ["目的", "实验目的", "实验要求"],
    "实验原理": ["原理", "实验原理", "基本原理", "理论基础"],
    "实验设备": ["设备", "仪器", "实验装置", "实验器材", "实验材料"],
    "实验步骤": ["步骤", "实验步骤", "操作步骤", "实验过程", "实验方法"],
    "实验数据": ["数据", "实验数据", "实验结果", "测量数据", "数据记录"],
    "数据处理": ["处理", "数据处理", "数据计算", "计算过程"],
    "结果分析": ["分析", "结果分析", "实验分析", "误差分析"],
    "实验感想": ["感想", "思考", "体会", "总结", "结论", "心得", "收获与体会"],
}


def identify_section_type(text: str) -> Optional[str]:
    """识别一段文字属于哪个章节类型"""
    text_lower = text.strip()
    for section_type, keywords in SECTION_KEYWORDS.items():
        for kw in keywords:
            if kw in text_lower and len(text_lower) <= 20:
                return section_type
    return None


def is_subjective_section(section_type: str) -> bool:
    """判断是否为主观题章节"""
    return section_type in ["实验感想", "结果分析"]


def extract_structure(extracted: dict) -> dict:
    """将提取的内容按章节结构化"""
    structure = {
        "title": "",
        "sections": [],
        "personal_info": {},
    }

    if extracted["source"] == "pdf":
        lines = extracted["text"].split("\n")
    else:
        lines = [p["text"] for p in extracted.get("paragraphs", [])]

    current_section = {"type": "header", "title": "报告头部", "content": []}

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # 尝试识别章节标题
        section_type = identify_section_type(line)

        if section_type:
            # 保存上一个 section
            if current_section["content"]:
                structure["sections"].append(current_section)
            current_section = {
                "type": section_type,
                "title": line,
                "content": []
            }
        else:
            current_section["content"].append(line)

    # 最后一个 section
    if current_section["content"]:
        structure["sections"].append(current_section)

    # 提取个人信息
    full_text = "\n".join(lines)
    name_match = re.search(r'姓\s*名[：:]\s*(\S+)', full_text)
    id_match = re.search(r'学\s*号[：:]\s*(\S+)', full_text)
    class_match = re.search(r'班\s*级[：:]\s*(\S+)', full_text)
    date_match = re.search(r'日\s*期[：:]\s*(\S+)', full_text)

    if name_match:
        structure["personal_info"]["name"] = name_match.group(1)
    if id_match:
        structure["personal_info"]["student_id"] = id_match.group(1)
    if class_match:
        structure["personal_info"]["class"] = class_match.group(1)
    if date_match:
        structure["personal_info"]["date"] = date_match.group(1)

    # 提取标题（通常在第一行）
    if lines:
        structure["title"] = lines[0].strip()

    return structure


# ============================================================================
# MD 模板生成
# ============================================================================

def generate_md_template(
    structure: dict,
    personal_info: dict,
    output_dir: str,
    image_paths: List[str] = None,
    data_paths: List[str] = None,
    subjective_answers: dict = None,
) -> str:
    """
    生成带占位符的 Markdown 模板

    Args:
        structure: 报告结构（extract_structure 的输出）
        personal_info: 用户个人信息 {name, student_id, class, date}
        output_dir: 输出目录
        image_paths: 用户提供的图片路径列表（示意图）
        data_paths: 用户提供的数据文件路径（Excel/CSV）
        subjective_answers: 主观题答案 {section_type: answer}

    Returns:
        生成的 MD 文件路径
    """
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    sources_dir = output_path / "sources"
    sources_dir.mkdir(exist_ok=True)

    md_lines = []
    info = personal_info or {}

    # ── 标题和个人信息 ──
    md_lines.append(f"# {structure.get('title', '实验报告')}")
    md_lines.append("")
    md_lines.append(f"**姓名**: {info.get('name', '[待填写]')}")
    md_lines.append(f"**学号**: {info.get('student_id', '[待填写]')}")
    md_lines.append(f"**班级**: {info.get('class', '[待填写]')}")
    md_lines.append(f"**日期**: {info.get('date', datetime.now().strftime('%Y-%m-%d'))}")
    md_lines.append("")
    md_lines.append("---")
    md_lines.append("")

    # ── 各章节 ──
    for section in structure.get("sections", []):
        sec_type = section["type"]
        sec_title = section["title"]
        sec_content = section["content"]

        md_lines.append(f"## {sec_title}")
        md_lines.append("")

        if is_subjective_section(sec_type):
            # 主观题章节
            if subjective_answers and sec_type in subjective_answers:
                md_lines.append(subjective_answers[sec_type])
            else:
                md_lines.append(f"[请填写你的{sec_title}]")
            md_lines.append("")
        else:
            # 客观内容：原文保留
            for line in sec_content:
                md_lines.append(line)
            md_lines.append("")

    # ── 图片引用（用户提供） ──
    if image_paths:
        md_lines.append("## 实验图片")
        md_lines.append("")
        for i, img_path in enumerate(image_paths):
            img_name = Path(img_path).name
            dest = sources_dir / img_name
            if os.path.exists(img_path):
                shutil.copy2(img_path, str(dest))
            md_lines.append(f"![图片{i+1}](sources/{img_name})")
            md_lines.append("")

    # ── 数据文件提示 ──
    if data_paths:
        md_lines.append("<!-- 数据文件: -->")
        for dp in data_paths:
            md_lines.append(f"<!-- - {dp} -->")
        md_lines.append("")

    # 写入文件
    md_content = "\n".join(md_lines)
    name = info.get('name', '报告')
    sid = info.get('student_id', '')
    cls = info.get('class', '')
    md_filename = f"{name}_{sid}_{cls}_实验报告.md" if (sid or cls) else "实验报告.md"
    md_path = output_path / md_filename
    md_path.write_text(md_content, encoding='utf-8')

    print(f"已生成 MD 模板: {md_path}")
    return str(md_path)


# ============================================================================
# MD → TEX → PDF 编译
# ============================================================================

def md_to_tex(md_path: str, output_dir: str = None) -> str:
    """将 Markdown 转为 LaTeX，复用 assignment 的转换器"""
    m2l = _import_assignment()

    md_content = Path(md_path).read_text(encoding='utf-8')
    md_dir = str(Path(md_path).parent)

    if output_dir is None:
        output_dir = md_dir

    if m2l:
        # 使用 assignment 的转换器
        title, info, submit_date = m2l.parse_markdown(md_content)
        latex_content = m2l.md_to_latex(md_content, title, info, submit_date)
    else:
        # 内建简易转换
        latex_content = _simple_md_to_tex(md_content)

    tex_path = Path(output_dir) / (Path(md_path).stem + ".tex")
    tex_path.write_text(latex_content, encoding='utf-8')
    print(f"已生成 TEX: {tex_path}")
    return str(tex_path)


def compile_to_pdf(tex_path: str) -> Optional[str]:
    """编译 TEX 为 PDF"""
    m2l = _import_assignment()

    if m2l:
        success = m2l.compile_pdf(tex_path)
    else:
        success = _simple_compile(tex_path)

    if success:
        pdf_path = str(Path(tex_path).with_suffix('.pdf'))
        print(f"已编译 PDF: {pdf_path}")
        return pdf_path
    else:
        print("PDF 编译失败")
        return None


def _simple_md_to_tex(md_content: str) -> str:
    """内建的简易 MD→TEX 转换（当 assignment 不可用时的备选）"""
    lines = md_content.split('\n')
    latex_lines = []
    in_list = False

    # 提取标题
    title = "实验报告"
    for line in lines:
        if line.startswith('# ') and not line.startswith('## '):
            title = line[2:].strip()
            break

    for line in lines:
        if not line.strip():
            if in_list:
                latex_lines.append(r'\end{itemize}')
                in_list = False
            latex_lines.append('')
            continue

        if line.startswith('### '):
            latex_lines.append(r'\subsubsection{' + line[4:].strip() + '}')
        elif line.startswith('## '):
            latex_lines.append(r'\subsection{' + line[3:].strip() + '}')
        elif line.startswith('# '):
            continue  # title handled in header
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            if not in_list:
                latex_lines.append(r'\begin{itemize}')
                in_list = True
            latex_lines.append(r'\item ' + line.strip()[2:])
        elif line.strip().startswith('!['):
            # 图片
            img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line.strip())
            if img_match:
                caption = img_match.group(1)
                img_path = img_match.group(2)
                latex_lines.append(r'\begin{figure}[H]')
                latex_lines.append(r'\centering')
                latex_lines.append(r'\includegraphics[width=0.8\textwidth]{' + img_path + '}')
                latex_lines.append(r'\caption{' + caption + '}')
                latex_lines.append(r'\end{figure}')
        elif line == '---':
            latex_lines.append(r'\newpage')
        else:
            converted = line
            converted = re.sub(r'\*\*(.+?)\*\*', r'\\textbf{\1}', converted)
            # 处理表格
            if '|' in converted and converted.strip().startswith('|'):
                continue  # 简易跳过表格行，需要更复杂的处理
            latex_lines.append(converted)

    if in_list:
        latex_lines.append(r'\end{itemize}')

    body = '\n'.join(latex_lines)

    return r"""\documentclass[12pt,a4paper]{ctexart}
\usepackage{amsmath}
\usepackage{amssymb}
\usepackage{geometry}
\usepackage{graphicx}
\usepackage{float}
\usepackage{booktabs}
\geometry{margin=2.5cm}

\title{""" + title + r"""}
\author{}
\date{}

\begin{document}
\maketitle

""" + body + r"""

\end{document}"""


def _simple_compile(tex_path: str) -> bool:
    """内建的简易 PDF 编译"""
    try:
        result = subprocess.run(
            ['pdflatex', '-interaction=nonstopmode', tex_path],
            capture_output=True, text=True,
            cwd=str(Path(tex_path).parent)
        )
        return result.returncode == 0
    except FileNotFoundError:
        print("错误: 未找到 pdflatex，请安装 MikTeX 或 TeX Live")
        return False


# ============================================================================
# CLI 入口
# ============================================================================

def main():
    import argparse
    parser = argparse.ArgumentParser(description='实验报告提取器')
    parser.add_argument('file', help='PDF/DOCX 文件路径')
    parser.add_argument('-o', '--output', help='输出目录')
    parser.add_argument('--name', default='', help='姓名')
    parser.add_argument('--student-id', default='', help='学号')
    parser.add_argument('--class', default='', help='班级')
    parser.add_argument('--date', default='', help='日期')
    parser.add_argument('--images', nargs='*', help='图片路径')
    parser.add_argument('--data', nargs='*', help='数据文件路径')
    parser.add_argument('--json', action='store_true', help='输出 JSON 结构')

    args = parser.parse_args()
    file_path = args.file

    # 提取内容
    ext = Path(file_path).suffix.lower()
    if ext == '.pdf':
        extracted = extract_pdf(file_path)
    elif ext in ('.docx', '.doc'):
        extracted = extract_docx(file_path)
    else:
        print(f"不支持的文件格式: {ext}")
        sys.exit(1)

    # 识别结构
    structure = extract_structure(extracted)

    if args.json:
        print(json.dumps(structure, ensure_ascii=False, indent=2))
        return

    # 输出目录
    if args.output:
        output_dir = args.output
    else:
        output_dir = str(Path(file_path).parent / "实验报告_改写")

    # 个人信息
    personal_info = {
        "name": args.name,
        "student_id": args.student_id,
        "class": args.__dict__.get('class', ''),
        "date": args.date,
    }

    # 生成 MD
    md_path = generate_md_template(
        structure=structure,
        personal_info=personal_info,
        output_dir=output_dir,
        image_paths=args.images,
        data_paths=args.data,
    )
    print(f"\nMD 模板已生成: {md_path}")


if __name__ == '__main__':
    main()
